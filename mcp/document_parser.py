"""
文档解析器 —— RAG 离线构造阶段的格式适配层。

职责：将上传的二进制文档解析为纯文本，交给知识库完成切块与向量化。
支持格式：
- .pdf  → pdfplumber 逐页抽取文本
- .docx → python-docx 抽取段落与表格文本
- .html → BeautifulSoup 剥离标签，保留正文与标题

解析结果统一为 [{"title": ..., "content": ...}]，
与 /knowledge/add 的文档格式一致，后续切片入库链路完全复用。
"""
import io
import logging
import re
from typing import Dict, List

logger = logging.getLogger(__name__)

# 各格式对应的后缀名集合（小写、含点）
SUPPORTED_BINARY_SUFFIXES = {".pdf", ".docx", ".html", ".htm"}


class DocumentParseError(Exception):
    """文档解析失败（格式损坏、加密、依赖缺失等）。"""


def parse_document(filename: str, raw: bytes) -> List[Dict[str, str]]:
    """
    按文件后缀分发到对应解析器。

    返回文档列表（多数格式为一篇；保留 List 以对齐知识库接口）。
    """
    suffix = _get_suffix(filename)
    title = _stem(filename)

    if suffix == ".pdf":
        content = _parse_pdf(raw)
    elif suffix == ".docx":
        content = _parse_docx(raw)
    elif suffix in (".html", ".htm"):
        content = _parse_html(raw, fallback_title=title)
        # HTML 优先用 <title> 作为文档标题
        extracted_title = _html_title(raw)
        if extracted_title:
            title = extracted_title
    else:
        raise DocumentParseError(f"不支持的文件格式: {suffix}")

    content = _clean_text(content)
    if not content:
        raise DocumentParseError(f"文件 {filename} 未解析出任何文本（可能是扫描件/图片型 PDF 或空文档）")

    return [{"title": title, "content": content}]


# ── 各格式解析实现 ────────────────────────────────────────────────────────────

def _parse_pdf(raw: bytes) -> str:
    """pdfplumber 逐页抽取文本，页间用换行分隔。"""
    try:
        import pdfplumber
    except ImportError:
        raise DocumentParseError("缺少依赖 pdfplumber，请先 pip install pdfplumber")

    pages: List[str] = []
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text)
    except DocumentParseError:
        raise
    except Exception as ex:
        msg = str(ex)
        if "password" in msg.lower() or "encrypt" in msg.lower():
            raise DocumentParseError("PDF 已加密，请先解密后再上传")
        raise DocumentParseError(f"PDF 解析失败（文件可能损坏或为扫描件）: {ex}")
    return "\n\n".join(pages)


def _parse_docx(raw: bytes) -> str:
    """python-docx 抽取段落与表格单元格文本。"""
    try:
        import docx
    except ImportError:
        raise DocumentParseError("缺少依赖 python-docx，请先 pip install python-docx")

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as ex:
        raise DocumentParseError(f"DOCX 解析失败（文件可能损坏）: {ex}")

    parts: List[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("；".join(cells))
    return "\n".join(parts)


def _parse_html(raw: bytes, fallback_title: str) -> str:
    """BeautifulSoup 剥离 script/style 后取正文文本。"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise DocumentParseError("缺少依赖 beautifulsoup4，请先 pip install beautifulsoup4")

    text = raw.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "noscript"]):
        tag.decompose()
    body = soup.get_text(separator="\n")
    # 压缩连续空行
    body = re.sub(r"\n{3,}", "\n\n", body)
    return body.strip() or fallback_title


def _html_title(raw: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw.decode("utf-8", errors="ignore"), "html.parser")
        return soup.title.get_text(strip=True) if soup.title else ""
    except Exception:
        return ""


# ── 工具函数 ─────────────────────────────────────────────────────────────────

def _get_suffix(filename: str) -> str:
    name = (filename or "").lower()
    return name[name.rfind("."):] if "." in name else ""


def _stem(filename: str) -> str:
    name = filename or "unknown"
    return name.rsplit(".", 1)[0] if "." in name else name


def _clean_text(text: str) -> str:
    """去除控制字符与多余空白，保留段落结构。"""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text or "")
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
